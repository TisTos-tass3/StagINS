from django.shortcuts import render, redirect, get_object_or_404
from .models import Stagiaire, Encadrant, Stage, Rapport
from .forms import StagiaireForm, StageForm, EncadrantForm, RapportForm

from django.http import JsonResponse, HttpResponse, Http404
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.core.exceptions import ValidationError
from django.utils import timezone
from django.db.models import Q
from django.db.models import Count, Avg, F
from django.db.models.functions import ExtractMonth
from django.contrib.auth import authenticate, login, logout, get_user_model
from django.contrib.auth.decorators import login_required 

import json
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
from collections import defaultdict
import os



from io import BytesIO
import tempfile


from docx.shared import Pt
import traceback

from rest_framework.decorators import permission_classes
from rest_framework.permissions import IsAuthenticated


from .permissions import IsAdmin, IsGestionnaireOrAdmin, CanEditStages, CanValidateRapports



from docx import Document
from docx.shared import Inches
from datetime import datetime
import pythoncom

# =======================================================================
# Fonctions de sérialisation 
# =======================================================================

def stagiaire_to_dict(stagiaire):
    if stagiaire is None:
        return None
    stages_count = stagiaire.stages_stagiaire.count() if hasattr(stagiaire, 'stages') else 0
    return {
        'id': stagiaire.id,
        'nom': stagiaire.nom,
        'prenom': stagiaire.prenom,
        'ecole': stagiaire.ecole,
        'specialite': stagiaire.specialite,
        'niveau_etude': stagiaire.niveau_etude,
        'email': stagiaire.email,
        'telephone': stagiaire.telephone,
        'matricule': stagiaire.matricule, 
        
        'stages_count': stages_count,
    }


def encadrant_to_dict(encadrant):
    if encadrant is None:
        return None
    stages_encadres = encadrant.stages_encadrant.count() if hasattr(encadrant, 'stages') else 0
    return {
        'id': encadrant.id,
        'nom': encadrant.nom,
        'prenom': encadrant.prenom,
        'institution': encadrant.institution,
        'nom_institution': encadrant.nom_institution,  
        'email': encadrant.email,
        'telephone': encadrant.telephone,
        'stages_encadres': stages_encadres,
    }




def stage_to_dict_light(stage):
    """Version légère"""
    return {
        'id': stage.id,
        'theme': stage.theme,
        'type_stage': stage.type_stage,
        'date_debut': stage.date_debut.isoformat() if stage.date_debut else None,
        'date_fin': stage.date_fin.isoformat() if stage.date_fin else None,
        'statut': stage.statut,
        'direction': stage.direction,
       
        'division': stage.division,
        
        'unite': stage.unite,
     
        'service': stage.service,
     
        'decision': stage.decision,
        'lettre_acceptation_url': stage.lettre_acceptation.url if stage.lettre_acceptation else None,
        'stagiaire_id': stage.stagiaire_id,
        'encadrant_id': stage.encadrant_id,
        'stagiaire': stagiaire_to_dict(stage.stagiaire),
        'encadrant': encadrant_to_dict(stage.encadrant),
        'rapports_count': stage.rapports.count() if hasattr(stage, 'rapports') else 0,
    }
def stage_to_dict(stage):
    """Version complète"""
    return stage_to_dict_light(stage)

 
def stagiaire_to_detail_dict(stagiaire):
    detail_dict = stagiaire_to_dict(stagiaire)
    
    stages = Stage.objects.filter(stagiaire=stagiaire).select_related('encadrant').order_by('-date_debut')
    
    detail_dict['stages'] = []
    
    for stage in stages:
        stage_rapport = Rapport.objects.filter(stage=stage).first()
        
        detail_dict['stages'].append({
            'id': stage.id,
            'theme': stage.theme,
            'date_debut': stage.date_debut.strftime('%Y-%m-%d'),
            'date_fin': stage.date_fin.strftime('%Y-%m-%d'),
            'statut': stage.statut,
            'type_stage': stage.type_stage,
            'direction': stage.direction,
            'division': stage.division,
            'unite': stage.unite,
            'service': stage.service,
            'encadrant_nom': f"{stage.encadrant.nom} {stage.encadrant.prenom}" if stage.encadrant else "Non assigné",
            'rapport': stage_rapport.etat if stage_rapport else None,
        })

    return detail_dict
  



def rapport_to_dict(rapport):
    return {
        'id': rapport.id,
        'date_depot': rapport.date_depot.isoformat() if rapport.date_depot else None,
        'etat': rapport.etat,
        'stage_id': rapport.stage_id,
        
        'fichier_url': rapport.fichier.url if rapport.fichier else None,
        'stage': stage_to_dict_light(rapport.stage), 
    }

# =======================================================================
# Vues d'interface utilisateur (HTML)
# =======================================================================

def home(request):
    stages = Stage.objects.all()
    return render(request, 'stages/home.html', {'stages': stages})

def add_stagiaire(request):
   
    return HttpResponse("Logique d'ajout de stagiaire via template (HTML)")

def add_stage(request):
    
    return HttpResponse("Logique d'ajout de stage via template (HTML)")


# =======================================================================
# Vues API Authentification 
# =======================================================================

@require_http_methods(["POST"])
@csrf_exempt 
def login_api(request):
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Format JSON invalide'}, status=400)
        
    username = data.get('username')
    password = data.get('password')
    
    user = authenticate(request, username=username, password=password)
    if user is not None:
        login(request, user)
        role = getattr(user, 'role', 'consultant')
       
        return JsonResponse({
            'success': True,
            'user': {
                'id': user.id,
                'username': user.username,
                'role': role,
                'email': user.email,
                'permissions': {
                    
                    'can_edit': role in ['admin', 'gestionnaire'],
                    'can_validate': role in ['admin', 'gestionnaire'],
                    'is_admin': user.is_superuser,
                }
            }
        })
    else:
        return JsonResponse({'success': False, 'error': 'Identifiants invalides'}, status=401)

@require_http_methods(["POST"])
def logout_api(request):
    logout(request)
    return JsonResponse({'success': True})

@require_http_methods(["GET"])
@csrf_exempt 
def current_user_api(request):
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Unauthorized'}, status=401)

    user = request.user
    role = getattr(user, 'role', 'consultant')
    return JsonResponse({
        'id': user.id,
        'username': user.username,
        'role': role,
        'email': user.email,
        'permissions': {
            'can_edit': role in ['admin', 'gestionnaire'],
            'can_validate': role in ['admin', 'gestionnaire'],
            'is_admin': user.is_superuser,
        }
    })


# =======================================================================
# Vues API Stagiaires
# =======================================================================

@require_http_methods(["GET"])
@permission_classes([IsAuthenticated])
def stagiaires_api(request):
  
    stagiaires = Stagiaire.objects.all().prefetch_related('stages_stagiaire').order_by('nom', 'prenom') 
    return JsonResponse([stagiaire_to_dict(s) for s in stagiaires], safe=False)


@csrf_exempt
@require_http_methods(["POST"])
@permission_classes([IsGestionnaireOrAdmin])
def stagiaire_create(request):
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Format JSON invalide'}, status=400)
    
    form = StagiaireForm(data)
    if form.is_valid():
        stagiaire = form.save()
        return JsonResponse(stagiaire_to_dict(stagiaire), status=201)
    return JsonResponse({'form_errors': form.errors}, status=400)

@csrf_exempt
@require_http_methods(["GET", "PUT", "DELETE"])
@permission_classes([IsGestionnaireOrAdmin])
def stagiaire_detail(request, pk):
    
    stagiaire = get_object_or_404(Stagiaire.objects.prefetch_related('stages_stagiaire'), pk=pk)

    if request.method == 'GET':
        return JsonResponse(stagiaire_to_dict(stagiaire))
    
    elif request.method == 'PUT':
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Format JSON invalide'}, status=400)
            
        form = StagiaireForm(data, instance=stagiaire)
        if form.is_valid():
            stagiaire = form.save()
            return JsonResponse(stagiaire_to_dict(stagiaire))
        return JsonResponse({'form_errors': form.errors}, status=400)
        
    elif request.method == 'DELETE':
       
        if stagiaire.stages_stagiaire.exists():
            return JsonResponse({'error': 'Impossible de supprimer un stagiaire avec des stages associés.'}, status=400)
        stagiaire.delete()
        return JsonResponse({'deleted': True})


@csrf_exempt
@require_http_methods(["GET"])
@permission_classes([IsAuthenticated])
def search_stagiaire_by_matricule(request):
    """ Recherche un stagiaire par son matricule et retourne son ID et ses infos de base. """
    matricule = request.GET.get('matricule', '').strip()

    if not matricule:
        return JsonResponse({"error": "Matricule manquant"}, status=400)

    try:
        stagiaire = Stagiaire.objects.get(matricule__iexact=matricule)
    except Stagiaire.DoesNotExist:
        return JsonResponse({"error": f"Stagiaire avec matricule {matricule} non trouvé."}, status=404)

    return JsonResponse(stagiaire_to_dict(stagiaire))

@csrf_exempt
@require_http_methods(["GET"])
@permission_classes([IsAuthenticated])
def stagiaire_detail_api(request, pk):
    """ API pour obtenir les détails d'un stagiaire et son historique de stages (via ID) """
    try:

        stagiaire = Stagiaire.objects.get(pk=pk)
    except Stagiaire.DoesNotExist:
        return JsonResponse({"error": "Stagiaire non trouvé"}, status=404)

    
    data = stagiaire_to_detail_dict(stagiaire) 
    return JsonResponse(data, safe=False)


# =======================================================================
# Vues API Encadrants 
# =======================================================================

@require_http_methods(["GET"])
@permission_classes([IsAuthenticated])
def encadrants_api(request):
   
    encadrants = Encadrant.objects.all().prefetch_related('stages_encadrant').order_by('nom', 'prenom')
    return JsonResponse([encadrant_to_dict(e) for e in encadrants], safe=False)

@csrf_exempt
@require_http_methods(["POST"])
@permission_classes([IsGestionnaireOrAdmin])
def encadrant_create(request):
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Format JSON invalide'}, status=400)
    
    
    if 'email' in data:
        if Encadrant.objects.filter(email=data['email']).exists():
            return JsonResponse({'email': ['Un encadrant avec cet email existe déjà.']}, status=400)
    
    if 'telephone' in data and data['telephone']:
        if Encadrant.objects.filter(telephone=data['telephone']).exists():
            return JsonResponse({'telephone': ['Un encadrant avec ce numéro de téléphone existe déjà.']}, status=400)
    
    form = EncadrantForm(data)
    if form.is_valid():
        encadrant = form.save()
        return JsonResponse(encadrant_to_dict(encadrant))
    return JsonResponse({'form_errors': form.errors}, status=400)
@csrf_exempt
@require_http_methods(["GET", "PUT", "DELETE"])
@permission_classes([IsGestionnaireOrAdmin])
def encadrant_detail(request, pk):
    try:
      
        encadrant = get_object_or_404(Encadrant.objects.prefetch_related('stages_encadrant'), pk=pk)
    
    except Encadrant.DoesNotExist:
        return JsonResponse({'error': 'Encadrant non trouvé'}, status=404)
    
    except Exception as e:
        return JsonResponse({'error': f'Erreur serveur: {str(e)}'}, status=500)

    if request.method == 'GET':
        try:
            return JsonResponse(encadrant_to_dict(encadrant))
        except Exception as e:
            return JsonResponse({'error': 'Erreur lors de la sérialisation des données'}, status=500)
    
    elif request.method == 'PUT':
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Format JSON invalide'}, status=400)
        
        try:
            
            if 'email' in data:
                existing = Encadrant.objects.filter(email=data['email']).exclude(pk=pk)
                if existing.exists():
                    return JsonResponse({'email': ['Un encadrant avec cet email existe déjà.']}, status=400)
            
           
            if 'telephone' in data and data['telephone']:
                existing = Encadrant.objects.filter(telephone=data['telephone']).exclude(pk=pk)
                if existing.exists():
                    return JsonResponse({'telephone': ['Un encadrant avec ce numéro de téléphone existe déjà.']}, status=400)
            
            form = EncadrantForm(data, instance=encadrant)
            if form.is_valid():
                encadrant = form.save()
                return JsonResponse(encadrant_to_dict(encadrant))
            return JsonResponse({'form_errors': form.errors}, status=400)
            
        except Exception as e:
            return JsonResponse({'error': 'Erreur lors de la mise à jour'}, status=500)

    elif request.method == 'DELETE':
        try:
           
            stages_count = encadrant.stages_encadrant.count()
            if stages_count > 0:
                return JsonResponse({
                    'error': f'Impossible de supprimer cet encadrant. Il est associé à {stages_count} stage(s) actif(s). Veuillez d\'abord réaffecter ces stages à un autre encadrant.'
                }, status=400)
            
            encadrant.delete()
            return JsonResponse({
                'deleted': True,
                'message': 'Encadrant supprimé avec succès'
            })
            
        except Exception as e:
            return JsonResponse({'error': 'Erreur lors de la suppression'}, status=500)

# =======================================================================
# Vues API Stages
# =======================================================================

@require_http_methods(["GET"])
@permission_classes([IsAuthenticated]) 
def stages_api(request):
    stages = Stage.objects.all().select_related('stagiaire', 'encadrant').order_by('-date_debut')
    
    statut = request.GET.get('statut')
    if statut:
        stages = stages.filter(statut=statut)
    
    return JsonResponse([stage_to_dict(s) for s in stages], safe=False)



@csrf_exempt
@require_http_methods(["POST"])
@permission_classes([CanEditStages]) 
def stage_create(request):
    try:
       
        stagiaire_id = request.POST.get('stagiaire')
        encadrant_id = request.POST.get('encadrant')

        if not stagiaire_id:
            return JsonResponse({'error': 'L\'ID du stagiaire est obligatoire pour créer le stage.'}, status=400)

        stage_data = {
            'theme': request.POST.get('theme'),
            'type_stage': request.POST.get('type_stage'),
            'date_debut': request.POST.get('date_debut'),
            'date_fin': request.POST.get('date_fin'),
            'stagiaire': stagiaire_id,
            'encadrant': encadrant_id,
            'unite': request.POST.get('unite'),
            'service': request.POST.get('service'),
            'decision': request.POST.get('decision'),
        }

       
        stage_form = StageForm(stage_data, request.FILES)
        if stage_form.is_valid():
            try:
                stage = stage_form.save()
               
                stage = Stage.objects.select_related('stagiaire', 'encadrant').get(pk=stage.pk)
                return JsonResponse(stage_to_dict(stage), status=201)
            except Exception as e:
                return JsonResponse({'error': f"Erreur lors de la sauvegarde: {e}"}, status=500)
        else:
            return JsonResponse({'form_errors': stage_form.errors}, status=400)

    except Exception as e:
        return JsonResponse({'error': f"Erreur lors de la création du stage: {str(e)}"}, status=500)


@csrf_exempt
@require_http_methods(["GET", "PUT", "DELETE"])
@permission_classes([CanEditStages]) 
def stage_detail(request, pk):
    stage = get_object_or_404(Stage.objects.select_related('stagiaire', 'encadrant'), pk=pk)

    if request.method == 'GET':
        return JsonResponse(stage_to_dict(stage))
    elif request.method == 'PUT':
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Format JSON invalide'}, status=400)
            
        form = StageForm(data, instance=stage)
        if form.is_valid():
            stage = form.save()
            stage = Stage.objects.select_related('stagiaire', 'encadrant').get(pk=stage.pk)
            return JsonResponse(stage_to_dict(stage))
        return JsonResponse({'form_errors': form.errors}, status=400)
    elif request.method == 'DELETE':
        if stage.rapports.exists():
            return JsonResponse({'error': 'Impossible de supprimer un stage avec un rapport associé.'}, status=400)
        stage.delete()
        return JsonResponse({'deleted': True})
    stage = get_object_or_404(Stage.objects.select_related('stagiaire', 'encadrant'), pk=pk)

    if request.method == 'GET':
        return JsonResponse(stage_to_dict(stage))
    elif request.method == 'PUT':
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Format JSON invalide'}, status=400)
            
        form = StageForm(data, instance=stage)
        if form.is_valid():
            stage = form.save()
            stage = Stage.objects.select_related('stagiaire', 'encadrant').get(pk=stage.pk)
            return JsonResponse(stage_to_dict(stage))
        return JsonResponse({'form_errors': form.errors}, status=400)
    elif request.method == 'DELETE':
        if stage.rapports.exists():
            return JsonResponse({'error': 'Impossible de supprimer un stage avec un rapport associé.'}, status=400)
        stage.delete()
        return JsonResponse({'deleted': True})


# =======================================================================
# Vues API Rapports 
# =======================================================================
@csrf_exempt
@require_http_methods(["GET"])
@permission_classes([IsAuthenticated])
def rapports_api(request):
    rapports = Rapport.objects.all().select_related('stage__stagiaire', 'stage__encadrant').order_by('-date_depot')
    
    query = request.GET.get('q', '').lower()
    etat_filter = request.GET.get('etat', '')
    annee_filter = request.GET.get('annee', '')

    if query:
        rapports = rapports.filter(
            Q(stage__theme__icontains=query) |
            Q(stage__stagiaire__nom__icontains=query) |
            Q(stage__stagiaire__prenom__icontains=query)
        )

    if etat_filter:
        rapports = rapports.filter(etat=etat_filter)

    if annee_filter and annee_filter.isdigit():
        try:
            annee = int(annee_filter)
            rapports = rapports.filter(date_depot__year=annee)
        except ValueError:
            pass

    return JsonResponse([rapport_to_dict(r) for r in rapports], safe=False)
@csrf_exempt
@require_http_methods(["POST"])
@permission_classes([IsAuthenticated])
def rapport_create(request):
    try:
        stage_id = request.POST.get("stage")
        fichier = request.FILES.get("fichier")

        if not stage_id:
            return JsonResponse({"stage": ["Le stage est obligatoire."]}, status=400)
        if not fichier:
            return JsonResponse({"fichier": ["Le fichier est obligatoire."]}, status=400)

        
        stage = Stage.objects.get(id=int(stage_id))

        
        rapport = Rapport.objects.create(stage=stage, fichier=fichier)

        return JsonResponse({
            "id": rapport.id,
            "stage": rapport.stage.id,
            "fichier_url": rapport.fichier.url if rapport.fichier else None
        }, status=201)

    except Stage.DoesNotExist:
        return JsonResponse({"stage": ["Stage introuvable."]}, status=404)
    except Exception as e:
        print("Erreur:", e)
        return JsonResponse({"error": str(e)}, status=500)

    print("DEBUG - Données POST reçues:", dict(request.POST))
    print("DEBUG - Fichiers reçus:", dict(request.FILES))
    
    
    data = request.POST.copy()
    
   
    if 'stage' in data:
        try:
            data['stage'] = int(data['stage'])
        except (ValueError, TypeError):
            return JsonResponse({'form_errors': {'stage': ['ID de stage invalide']}}, status=400)
    
    form = RapportForm(data, request.FILES)
    
    print("DEBUG - Formulaire valide:", form.is_valid())
    print("DEBUG - Erreurs du formulaire:", form.errors if not form.is_valid() else "Aucune erreur")
    
    if form.is_valid():
        try:
           
            from .business_rules import BusinessRules
            
           
            BusinessRules.valider_creation_rapport(
                form.cleaned_data, 
                request.FILES.get('fichier')
            )
            
           
            rapport = form.save()
            rapport = Rapport.objects.select_related('stage__stagiaire', 'stage__encadrant').get(pk=rapport.pk)
            return JsonResponse(rapport_to_dict(rapport), status=201)
            
        except ValidationError as e:
          
            print("DEBUG - Erreur de validation métier:", e.message_dict)
            return JsonResponse({'form_errors': e.message_dict}, status=400)
        except Exception as e:
            print("DEBUG - Erreur lors de la sauvegarde:", str(e))
            return JsonResponse({'error': f"Erreur lors de la sauvegarde: {str(e)}"}, status=500)
    
    print("DEBUG - Erreurs de formulaire finales:", form.errors)
    return JsonResponse({'form_errors': form.errors}, status=400)
    print("DEBUG - Données POST reçues:", dict(request.POST))
    print("DEBUG - Fichiers reçus:", dict(request.FILES))
    
    try:
        stage_id = request.POST.get('stage')
        fichier = request.FILES.get('fichier')
        
        if not stage_id:
            return JsonResponse({'form_errors': {'stage': ['Stage est requis']}}, status=400)
        
        try:
            stage = Stage.objects.get(id=int(stage_id))
        except (Stage.DoesNotExist, ValueError):
            return JsonResponse({'form_errors': {'stage': ['Stage invalide']}}, status=400)
        
      
        rapport = Rapport(
            stage=stage,
            fichier=fichier,
            etat='En attente'  
        )
        
       
        from .business_rules import BusinessRules
        BusinessRules.valider_creation_rapport(
            {'stage': stage, 'fichier': fichier},
            fichier
        )
        
        rapport.save()
        
        
        rapport = Rapport.objects.select_related('stage__stagiaire', 'stage__encadrant').get(pk=rapport.pk)
        return JsonResponse(rapport_to_dict(rapport), status=201)
        
    except ValidationError as e:
        return JsonResponse({'form_errors': e.message_dict}, status=400)
    except Exception as e:
        return JsonResponse({'error': f"Erreur lors de la sauvegarde: {str(e)}"}, status=500)
    
    form = RapportForm(request.POST, request.FILES)
    
    if form.is_valid():
        try:
          
            from .business_rules import BusinessRules
            
          
            BusinessRules.valider_creation_rapport(
                form.cleaned_data, 
                request.FILES.get('fichier')
            )
            
           
            rapport = form.save()
            rapport = Rapport.objects.select_related('stage__stagiaire', 'stage__encadrant').get(pk=rapport.pk)
            return JsonResponse(rapport_to_dict(rapport), status=201)
            
        except ValidationError as e:
           
            return JsonResponse({'form_errors': e.message_dict}, status=400)
        except Exception as e:
            return JsonResponse({'error': f"Erreur lors de la sauvegarde: {str(e)}"}, status=500)
    
    return JsonResponse({'form_errors': form.errors}, status=400)
@csrf_exempt
@require_http_methods(["GET", "PUT", "DELETE"]) 
@permission_classes([IsGestionnaireOrAdmin])
def rapport_detail(request, pk):
    rapport = get_object_or_404(Rapport.objects.select_related('stage__stagiaire', 'stage__encadrant'), pk=pk)

    if request.method == 'GET':
        return JsonResponse(rapport_to_dict(rapport))
    
    elif request.method == 'PUT':
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Format JSON invalide'}, status=400)
            
      
    elif request.method == 'DELETE':
        try:
            from .business_rules import BusinessRules
            peut_supprimer, message = BusinessRules.peut_supprimer_rapport(rapport)
            
            if not peut_supprimer:
                return JsonResponse({'error': message}, status=400)
                
            if rapport.fichier and os.path.exists(rapport.fichier.path):
                os.remove(rapport.fichier.path)
                
            rapport.delete()
            return JsonResponse({'deleted': True})
            
        except Exception as e:
            return JsonResponse({'error': f'Erreur lors de la suppression: {str(e)}'}, status=500)
@csrf_exempt
@require_http_methods(["POST"])
@permission_classes([CanValidateRapports]) 
def rapport_valider(request, pk):
    rapport = get_object_or_404(Rapport.objects.select_related('stage__stagiaire', 'stage__encadrant'), pk=pk)
    
    if rapport.etat == 'Validé':
        return JsonResponse({"error": "Déjà validé."}, status=400)
        
    rapport.etat = 'Validé'
    rapport.save()
    
    
    stage = rapport.stage
    if stage.statut != 'Validé':
        stage.statut = 'Validé'
        stage.save()
        
    return JsonResponse(rapport_to_dict(rapport))

@csrf_exempt
@require_http_methods(["POST"])
@permission_classes([CanValidateRapports]) 
def rapport_archiver(request, pk):
    rapport = get_object_or_404(Rapport.objects.select_related('stage__stagiaire', 'stage__encadrant'), pk=pk)
    
    if rapport.etat != 'Validé':
        return JsonResponse({"error": "Un rapport ne peut être archivé que s'il est Validé."}, status=400)
        
    rapport.etat = 'Archivé'
    rapport.save()
    return JsonResponse(rapport_to_dict(rapport))

@csrf_exempt
@require_http_methods(["GET"])
@permission_classes([IsAuthenticated])
def rapport_download(request, pk):
    rapport = get_object_or_404(Rapport, pk=pk)
    
    if not rapport.fichier:
        return JsonResponse({"error": "Fichier non trouvé"}, status=404)
        
   
    file_path = rapport.fichier.path
    if not os.path.exists(file_path):
        raise Http404("Fichier sur le disque introuvable.")

    with open(file_path, 'rb') as fh:
        response = HttpResponse(fh.read(), content_type="application/pdf")
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
        return response
    



# =======================================================================
# Vues API Dashboard 
# =======================================================================

def get_monthly_stage_count():
    """Compte le nombre de stages qui ont commencé chaque mois de l'année En cours."""
    current_year = date.today().year
    
    monthly_counts_qs = Stage.objects.filter(
        date_debut__year=current_year
    ).annotate(
        month=ExtractMonth('date_debut')
    ).values('month').annotate(
        count=Count('id')
    ).order_by('month')
    
    monthly_data = {i: 0 for i in range(1, 13)}
    for item in monthly_counts_qs:
        monthly_data[item['month']] = item['count']
        
    month_names = ["Jan", "Fév", "Mar", "Avr", "Mai", "Juin", "Juil", "Aoû", "Sep", "Oct", "Nov", "Déc"]
    
    formatted_data = [
        {'name': month_names[i-1], 'Stages': monthly_data[i]}
        for i in range(1, 13)
    ]
    
    return formatted_data


def get_dashboard_data():
    """Collecte toutes les statistiques pour le dashboard."""
    total_stagiaires = Stagiaire.objects.count()
    total_encadrants = Encadrant.objects.count()
    
    stages_counts = Stage.objects.values('statut').annotate(count=Count('statut'))
    stages_dict = {item['statut']: item['count'] for item in stages_counts}
    total_stages = Stage.objects.count()
    
    rapports_counts = Rapport.objects.values('etat').annotate(count=Count('etat'))
    rapports_dict = {item['etat']: item['count'] for item in rapports_counts}
    total_rapports = Rapport.objects.count()

    today = date.today()
    stages_a_verifier = Stage.objects.filter(
        statut__in=['En cours', 'En attente'], 
        date_fin__lt=today
    ).count() 
    
    seven_days_later = today + relativedelta(days=+7)
    stages_bientot_finis = Stage.objects.filter(
        statut='En cours', 
        date_fin__range=[today, seven_days_later]
    ).count()
    
    stage_status_data = [
        {'name': statut, 'value': stages_dict.get(statut, 0)}
        for statut in stages_dict.keys()
    ]
    
    monthly_stages_data = get_monthly_stage_count()
    
    return {
        'total_stagiaires': total_stagiaires,
        'total_encadrants': total_encadrants,
        'stages_encours': stages_dict.get('En cours', 0),
        'stages_valides': stages_dict.get('Validé', 0),
        'stages_termines': stages_dict.get('Terminé', 0),
        'total_stages': total_stages,
        'rapports_en_attente': rapports_dict.get('En attente', 0),
        'rapports_valides': rapports_dict.get('Validé', 0),
        'rapports_archives': rapports_dict.get('Archivé', 0),
        'total_rapports': total_rapports,
        'stages_retard_non_valides': stages_a_verifier,
        'stages_bientot_finis': stages_bientot_finis,
        'stages_by_status': stage_status_data, 
        'monthly_stages': monthly_stages_data, 
    }

@require_http_methods(["GET"])
@permission_classes([IsAuthenticated])
def dashboard_api(request):
    """API endpoint pour les données du dashboard."""
    data = get_dashboard_data()
    return JsonResponse(data)


# =======================================================================
# VUES POUR LA RECHERCHE ET LE DÉTAIL DU DOSSIER STAGIAIRE
# =======================================================================

def stagiaire_to_detail_dict(stagiaire):
    """
    Sérialise le stagiaire avec tous ses stages et rapports associés.
    (Fonction déplacée au niveau global)
    """
    detail_dict = stagiaire_to_dict(stagiaire)
    
    stages = Stage.objects.filter(stagiaire=stagiaire).select_related('encadrant').order_by('-date_debut')
    
    detail_dict['stages'] = []
    
    for stage in stages:
        stage_rapport = Rapport.objects.filter(stage=stage).first() 
        
        encadrant_type = 'Interne' if getattr(stage.encadrant, 'interne', False) else 'Externe'
        encadrant_nom = f"{stage.encadrant.nom} {stage.encadrant.prenom} ({encadrant_type})"
        
        
        rapport_dict = None
        if stage_rapport:
            rapport_dict = {
                'statut': stage_rapport.etat,
                'date_depot': stage_rapport.date_depot.strftime('%Y-%m-%d') if stage_rapport.date_depot else None,
               
                'download_url': f"/api/rapport/{stage_rapport.id}/download/", 
            }
        
        detail_dict['stages'].append({
            'id': stage.id,
            'theme': stage.theme,
            'date_debut': stage.date_debut.strftime('%Y-%m-%d'),
            'date_fin': stage.date_fin.strftime('%Y-%m-%d'),
            'statut': stage.statut,
            'type_stage': stage.type_stage,
            'unite': stage.unite,
            'service': stage.service,
            'encadrant_nom': encadrant_nom,
            'rapport': rapport_dict,
        })

    return detail_dict

def stagiaire_to_detail_dict(stagiaire):
    """
    Sérialise le stagiaire avec tous ses stages et rapports associés.
    """
    detail_dict = stagiaire_to_dict(stagiaire)
    
    stages = Stage.objects.filter(stagiaire=stagiaire).select_related('encadrant').order_by('-date_debut')
    
    detail_dict['stages'] = []
    
    for stage in stages:
        stage_rapport = Rapport.objects.filter(stage=stage).first() 
        
        encadrant_nom = f"{stage.encadrant.nom} {stage.encadrant.prenom}" if stage.encadrant else "Non assigné"
        
        rapport_dict = None
        if stage_rapport:
            rapport_dict = {
                'statut': stage_rapport.etat,
                'date_depot': stage_rapport.date_depot.strftime('%Y-%m-%d') if stage_rapport.date_depot else None,
                'download_url': f"/rapports/api/{stage_rapport.id}/download/", 
            }
        
        detail_dict['stages'].append({
            'id': stage.id,
            'theme': stage.theme,
            'date_debut': stage.date_debut.strftime('%Y-%m-%d'),
            'date_fin': stage.date_fin.strftime('%Y-%m-%d'),
            'statut': stage.statut,
            'type_stage': stage.type_stage,
         
            'direction': stage.direction, 

            'division': stage.division,  
            
           
            'unite': stage.unite,  
           
            'service': stage.service, 
           
            'encadrant_nom': encadrant_nom,
            'rapport': rapport_dict,
        })

    return detail_dict

@csrf_exempt
@require_http_methods(["GET"])
@permission_classes([IsAuthenticated])
def stagiaire_detail_api(request, pk):
    """ API pour obtenir les détails d'un stagiaire et son historique de stages (via ID) """
    try:
        stagiaire = Stagiaire.objects.get(pk=pk)
    except Stagiaire.DoesNotExist:
        return JsonResponse({"error": "Stagiaire non trouvé"}, status=404)

    data = stagiaire_to_detail_dict(stagiaire)
    return JsonResponse(data, safe=False)

# =======================================================================
# VUES POUR LA GÉNÉRATION D'ATTESTATIONS DE FIN DE STAGE
# =======================================================================

def get_direction_full_name(direction_code, unite=None):
    """Retourne le nom complet de la direction avec formatage spécial pour BCR"""
    direction_mapping = {
        'DGE': 'Direction Générale des Études ',
        'DF': 'Direction des Finances', 
        'DRH': 'Direction des Ressources Humaines',
        'DSI': 'Direction des Systèmes d\'Information',
        'BCR': 'Bureau Central du Recensement',
    }
    
    nom_direction = direction_mapping.get(direction_code, direction_code or "Non spécifiée")
    
    
    if direction_code == 'BCR' and unite:
        return f"{nom_direction}({unite})"
    
    return nom_direction

def get_direction_avec_article(direction_code, unite=None):
    """Retourne le nom de la direction avec l'article approprié ('du' pour BCR, 'de la' pour les autres)"""
    nom_direction = get_direction_full_name(direction_code, unite)
    
    if direction_code == 'BCR':
        return f"du {nom_direction}"
    else:
        return f"de la {nom_direction}"

def calculer_duree_mois(date_debut, date_fin):
    """Calcule la durée en mois entre deux dates"""
    if not date_debut or not date_fin:
        return "Non spécifiée"
    
    delta = relativedelta(date_fin, date_debut)
    mois = delta.years * 12 + delta.months
    if delta.days > 0:
        mois += 1
    
    return f"{mois} mois"

def remplacer_placeholders_robuste(doc, contexte):
    """Remplace les placeholders dans le document Word"""
    
    champs_en_gras = [
        'nom_complet', 'diplome', 'niveau_etude', 
        'duree_mois', 'date_debut_lettres', 'date_fin_lettres',
        'signataire','unite','service'
    ]

    def traiter_texte_avec_gras(texte):
        """Traite le texte pour appliquer le gras uniquement sur les noms des directions/unités/services"""
       
        champs_direction = ['direction', 'unite', 'service']
        
        for champ in champs_direction:
            if champ in contexte and contexte[champ]:
                valeur = str(contexte[champ])
                article = ""
                nom = valeur
                
               
                if valeur.startswith("du "):
                    article = "du "
                    nom = valeur[3:]  
                elif valeur.startswith("de la "):
                    article = "de la "
                    nom = valeur[6:] 
                
                
                if valeur in texte:
                    texte = texte.replace(valeur, f"{article}**{nom}**")
        
       
        for champ in champs_en_gras:
            valeur = str(contexte.get(champ, ''))
            if valeur and valeur in texte:
                texte = texte.replace(valeur, f"**{valeur}**")
        
        return texte

    def appliquer_formatage_paragraph(paragraph):
        """Applique le formatage à un paragraphe"""
        original_text = paragraph.text
        if not original_text.strip():
            return
        
       
        nouveau_texte = original_text
        for key, value in contexte.items():
            placeholder = f"{{{key}}}"
            if placeholder in nouveau_texte:
                nouveau_texte = nouveau_texte.replace(placeholder, str(value))
        
       
        nouveau_texte = traiter_texte_avec_gras(nouveau_texte)
        
       
        if nouveau_texte != original_text:
            paragraph.clear()
            
          
            parts = []
            current_text = nouveau_texte
            
            while "**" in current_text:
                before, rest = current_text.split("**", 1)
                if "**" in rest:
                    bold_text, after = rest.split("**", 1)
                    
                    if before:
                        parts.append({'text': before, 'bold': False})
                    
                    parts.append({'text': bold_text, 'bold': True})
                    current_text = after
                else:
                    if before:
                        parts.append({'text': before, 'bold': False})
                    parts.append({'text': rest, 'bold': True})
                    current_text = ""
                    break
            
            if current_text:
                parts.append({'text': current_text, 'bold': False})
            
            for part in parts:
                run = paragraph.add_run(part['text'])
                run.font.name = 'Century Gothic'
                run.font.size = Pt(14)
                run.bold = part['bold']

   
    for paragraph in doc.paragraphs:
        appliquer_formatage_paragraph(paragraph)
    
   
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    appliquer_formatage_paragraph(paragraph)
def appliquer_style_global(doc):
    """Applique la police Century Gothic taille 14 à tout le document"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Century Gothic'
    font.size = Pt(14)
    
   
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Century Gothic'
            run.font.size = Pt(14)
    
   
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Century Gothic'
                        run.font.size = Pt(14)

def generer_document_avec_template(stage, contexte, format_output):
    """Génère le document en utilisant le template Word"""
    try:
       
        if stage.type_stage == 'Academique':
            template_name = 'template_academique.docx'
        else:
            template_name = 'template_professionnel.docx'
        
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        possible_paths = [
            os.path.join(base_dir, 'templates', template_name),
            os.path.join(os.path.dirname(__file__), 'templates', template_name),
            os.path.join(base_dir, 'stages', 'templates', template_name),
        ]
        
        template_path = None
        for path in possible_paths:
            if os.path.exists(path):
                template_path = path
                break
        
        if not template_path:
            raise FileNotFoundError(f"Aucun template trouvé pour le type de stage '{stage.type_stage}'")
        
        
        doc = Document(template_path)
        
       
        remplacer_placeholders_robuste(doc, contexte)
        
      
        appliquer_style_global(doc)
        
       
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        
        filename = f"attestation_{stage.stagiaire.nom}_{stage.stagiaire.prenom}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
        
    except Exception as e:
        print(f"Erreur dans generer_document_avec_template: {str(e)}")
        raise

@csrf_exempt
@require_http_methods(["POST"])
@permission_classes([IsAuthenticated])

def generer_attestation(request, pk):
    """Génère une attestation de fin de stage pour un stage validé"""
    try:
     
        stage = get_object_or_404(Stage.objects.select_related('stagiaire', 'encadrant'), pk=pk)
        stagiaire = get_object_or_404(Stagiaire, pk=stage.stagiaire.pk)
        
      
        if stage.statut != 'Validé':
            return JsonResponse({'error': 'Seuls les stages validés peuvent générer une attestation.'}, status=400)
        
       
        data = json.loads(request.body)
        signataire = data.get('signataire', 'Tassiou ABOUBACAR')
        fonction_signataire = data.get('fonction_signataire', 'Directeur des Ressources Humaines')
        format_output = data.get('format', 'docx')
        
        direction_avec_article = get_direction_avec_article(stage.direction, stage.unite)
        
        duree_mois = calculer_duree_mois(stage.date_debut, stage.date_fin)
        
        def format_date_lettres(date_obj):
            mois = [
                'janvier', 'février', 'mars', 'avril', 'mai', 'juin',
                'juillet', 'août', 'septembre', 'octobre', 'novembre', 'décembre'
            ]
            return f"{date_obj.day} {mois[date_obj.month - 1]} {date_obj.year}"
        
        contexte = {
            'nom_complet': f"{stage.stagiaire.nom} {stage.stagiaire.prenom.upper()}",
            'diplome': f"{stagiaire.niveau_etude} ({stagiaire.specialite})" or "Non spécifié",
            'niveau_etude': stagiaire.niveau_etude or "Non spécifié",
            'theme': stage.theme,
            'duree_mois': duree_mois,
            'date_debut_lettres': format_date_lettres(stage.date_debut),
            'date_fin_lettres': format_date_lettres(stage.date_fin),
            'direction': direction_avec_article, 
            'unite': stage.unite or "Non spécifiée",
            'service': stage.service or "Non spécifié", 
            'encadrant': f"{stage.encadrant.prenom} {stage.encadrant.nom}" if stage.encadrant else "Non assigné",
            'signataire': signataire,
            'fonction_signataire': fonction_signataire,
            'date_signature': datetime.now().strftime('%d/%m/%Y'),
        }
        
        contexte['unite_complete'] = contexte['direction']
        
        print("=== CONTEXTE POUR GÉNÉRATION ATTESTATION ===")
        for key, value in contexte.items():
            print(f"{key}: {value}")
        print("=== FIN DU CONTEXTE ===")
        
        result = generer_document_avec_template(stage, contexte, format_output)
        return result
            
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Données JSON invalides'}, status=400)
    except Exception as e:
        print(f"Erreur lors de la génération: {str(e)}")
        return JsonResponse({'error': f'Erreur lors de la génération: {str(e)}'}, status=500)